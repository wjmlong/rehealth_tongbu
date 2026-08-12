from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ReHealth_APP_全景说明与技术架构_2026-08-12.md"
OUTPUT = ROOT / "docs" / "ReHealth_APP_全景说明与技术架构_2026-08-12.docx"

NAVY = "19324A"
BLUE = "2E74B5"
TEAL = "138A83"
MINT = "E8F5F2"
PALE_BLUE = "EAF1F8"
PALE_GOLD = "FFF4D9"
PALE_RED = "FBE9E7"
INK = "1F2933"
MUTED = "5D6974"
LINE = "D8E0E7"
WHITE = "FFFFFF"
RED = "B23A32"
GREEN = "18785F"
GOLD = "B47A13"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int], indent=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepNext"))
    if value and element is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and element is not None:
        p_pr.remove(element)


def set_keep_lines(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepLines"))
    if value and element is None:
        p_pr.append(OxmlElement("w:keepLines"))


def set_page_break_before(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:pageBreakBefore"))
    if value and element is None:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, color=LINE, size=8, space=4, side="bottom") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.84)
    section.right_margin = Inches(0.84)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 15, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name.startswith("Heading")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3.5)
        style.paragraph_format.line_spacing = 1.15

    for name, color, fill in (
        ("Lead Callout", NAVY, PALE_BLUE),
        ("Warning Callout", RED, PALE_RED),
        ("Success Callout", GREEN, MINT),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.2)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.left_indent = Pt(8)
        style.paragraph_format.right_indent = Pt(8)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15
        style._rehealth_fill = fill  # task-local marker


def add_running_furniture(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        left = p.add_run("REHEALTH AI  |  睿禾精灵 APP 全景报告")
        set_run_font(left, size=8.5, color=MUTED, bold=True)
        p.add_run("\t")
        right = p.add_run("管理层阅览版 · 2026-08-12")
        set_run_font(right, size=8.5, color=MUTED)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.75), 2)
        set_paragraph_border(p, color=LINE, size=5, space=3)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(2)
        r = fp.add_run("睿禾健康 · 内部材料     ")
        set_run_font(r, size=8, color=MUTED)
        add_field(fp, "PAGE")
        r2 = fp.add_run(" / ")
        set_run_font(r2, size=8, color=MUTED)
        add_field(fp, "NUMPAGES")


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].clear()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REHEALTH AI / 睿禾健康")
    set_run_font(r, size=11, color=TEAL, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(22)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("睿禾精灵 APP")
    set_run_font(r, size=30, color=NAVY, bold=True)
    r.add_break()
    r2 = title.add_run("全景说明与技术架构")
    set_run_font(r2, size=25, color=BLUE, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(20)
    rs = sub.add_run("产品、功能、技术栈、前后端、数据、AI、部署与发布状态")
    set_run_font(rs, size=13, color=MUTED)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.left_indent = Inches(1.15)
    status.paragraph_format.right_indent = Inches(1.15)
    status.paragraph_format.space_before = Pt(8)
    status.paragraph_format.space_after = Pt(26)
    shade_paragraph(status, PALE_RED)
    sr = status.add_run("当前结论：MVP 主体已形成 · 正式发布仍处于 BLOCKED 状态")
    set_run_font(sr, size=11, color=RED, bold=True)

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    set_table_width(meta, [2100, 7260], indent=120)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("汇报对象", "公司管理层、产品负责人、技术负责人"),
        ("客户端", "睿禾精灵 Android APP（com.rehealth.genie）"),
        ("待发布版本", "1.0.0（versionCode 1）"),
        ("现状基准", "2026-08-12 代码仓与权威状态文档"),
        ("文档版本", "V1.0 · 管理层阅览版"),
    ]
    for row, (label, value) in zip(meta.rows, data):
        set_table_row_cant_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            set_cell_margins(cell, 120, 140, 120, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], PALE_BLUE)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        rr = p0.add_run(label)
        set_run_font(rr, size=9.5, color=BLUE, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        rr = p1.add_run(value)
        set_run_font(rr, size=9.5, color=INK)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    nr = note.add_run("内部汇报材料 · 健康管理参考系统 · 不替代医疗诊断")
    set_run_font(nr, size=9, color=MUTED, italic=True)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    title = p.add_run("目录")
    set_run_font(title, size=16, color=BLUE, bold=True)
    set_paragraph_border(p, color=BLUE, size=10, space=5)
    toc = doc.add_paragraph()
    toc.paragraph_format.space_before = Pt(8)
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u')
    hint = doc.add_paragraph("提示：在 Word 中按 Ctrl+A 后按 F9，可更新目录和总页数。")
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint.paragraph_format.space_before = Pt(8)
    for run in hint.runs:
        set_run_font(run, size=8.5, color=MUTED, italic=True)
    doc.add_page_break()


def strip_inline(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text).replace("**", "")


def add_rich_text(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|https?://[^\s)]+)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=10.2, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=10.2, color=INK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.2, color=NAVY)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "F1F4F7")
            run._r.get_or_add_rPr().append(shade)
        else:
            add_hyperlink(paragraph, token, token)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=10.2, color=INK)


def add_business_summary(doc: Document) -> None:
    h = doc.add_paragraph("一页式管理摘要", style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    set_paragraph_border(h, color=BLUE, size=10, space=5)
    lead = doc.add_paragraph(style="Lead Callout")
    shade_paragraph(lead, PALE_BLUE)
    lead.paragraph_format.space_before = Pt(6)
    r = lead.add_run(
        "睿禾精灵已形成可演示、可联调的 Android 健康管理 MVP 主体：正式设备入口已收敛为 HBand MT116 与云米云端，"
        "数据采用本地优先和离线队列，云端按硬件遥测、业务编排和模型推理分层。当前仍需关闭真实设备、生产模型、"
        "签名包全链路和生产容灾等发布门禁。"
    )
    set_run_font(r, size=11, color=NAVY, bold=True)

    cards = doc.add_table(rows=2, cols=3)
    cards.style = "Table Grid"
    set_table_width(cards, [3120, 3120, 3120], indent=120)
    content = [
        ("2 类", "正式设备入口", "HBand BLE / 云米 IMEI"),
        ("5 个", "主导航模块", "首页 / 数据 / 归因 / 模型 / 我的"),
        ("3 个", "数据存储域", "Room / MySQL / TimescaleDB"),
        ("16 项", "CVD 风险输入", "含来源与数据质量"),
        ("22 张", "APP 本地表", "Room schema v16"),
        ("BLOCKED", "发布结论", "主体完成，生产门禁未关"),
    ]
    for row in cards.rows:
        set_table_row_cant_split(row)
    for i, (metric, label, detail) in enumerate(content):
        cell = cards.cell(i // 3, i % 3)
        set_cell_shading(cell, MINT if i < 5 else PALE_RED)
        set_cell_margins(cell, 150, 130, 150, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(metric)
        set_run_font(r, size=15, color=TEAL if i < 5 else RED, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(label)
        set_run_font(r2, size=9.5, color=INK, bold=True)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(detail)
        set_run_font(r3, size=8, color=MUTED)

    doc.add_paragraph("领导应重点关注", style="Heading 2")
    points = [
        "产品不是单一手环看板，而是设备采集、风险评估、干预和反馈闭环。",
        "正式能力与 Debug 工程能力已分离，Release 不包含 Mock、MRD/RWFit 工程入口。",
        "模型页和部分“我的”页入口仍有展示/占位内容，不能作为已交付能力宣传。",
        "当前最适合受控试点和真机验收，不应描述为已经全面正式上线。",
    ]
    for item in points:
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, item)

    doc.add_paragraph("管理决策建议", style="Heading 2")
    decision = doc.add_table(rows=4, cols=3)
    decision.style = "Table Grid"
    set_table_width(decision, [1500, 3060, 4800], indent=120)
    data = [
        ("优先级", "本阶段目标", "建议决策"),
        ("P0", "关闭发布阻塞", "集中完成 MT116 真机长稳、正式模型和签名包全链路验收"),
        ("P1", "补齐产品闭环", "完成 RHI 云端落库、隐私/导出/删除以及保险授权页面"),
        ("P2", "试点规模化", "再投入容量、容灾、模型漂移监测和多设备/IoT 扩展"),
    ]
    for r_idx, row_data in enumerate(data):
        row = decision.rows[r_idx]
        set_table_row_cant_split(row)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_margins(cell, 100, 120, 100, 120)
            set_cell_shading(cell, NAVY if r_idx == 0 else WHITE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            rr = p.add_run(value)
            set_run_font(rr, size=9, color=WHITE if r_idx == 0 else INK, bold=(r_idx == 0 or c_idx == 0))
    set_repeat_table_header(decision.rows[0])
    doc.add_page_break()


def add_architecture_visual(doc: Document) -> None:
    h = doc.add_paragraph("系统架构总览", style="Heading 2")
    set_keep_with_next(h)
    rows = [
        ("设备入口", "HBand MT116 BLE   |   云米 S8/S9/GS20/GS17/A67/K9L 云端"),
        ("移动端", "Compose UI → Provider 路由 → Room v16 → durable queue → WorkManager"),
        ("统一入口", "Edge / Gateway / HTTPS / JWT / 路由与安全头"),
        ("服务层", "Device Service（遥测）   |   JeecgBoot（业务）   |   model-service / PIAS（模型）"),
        ("数据与事件", "TimescaleDB + Outbox + Kafka   |   MySQL software_db   |   Redis / Nacos"),
        ("运营监控", "JeecgBoot Vue3 管理前端   |   Prometheus   |   Grafana"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1700, 7660], indent=120)
    fills = [PALE_GOLD, MINT, PALE_BLUE, "EDF2F7", "F4F6F8", "F7F8FA"]
    for idx, (label, value) in enumerate(rows):
        set_table_row_cant_split(table.rows[idx])
        for cell in table.rows[idx].cells:
            set_cell_margins(cell, 125, 140, 125, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, fills[idx])
        p0 = table.rows[idx].cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        set_run_font(r0, size=9.3, color=NAVY, bold=True)
        p1 = table.rows[idx].cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=9.1, color=INK)


def parse_markdown(doc: Document, source: str) -> None:
    lines = source.splitlines()
    idx = 0
    list_counter_active = False
    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()
        if idx < 10 and (not stripped or stripped.startswith("# ") or stripped.startswith(">") or stripped == "---"):
            idx += 1
            continue
        if not stripped or stripped == "---":
            idx += 1
            list_counter_active = False
            continue
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            block = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                block.append(lines[idx].rstrip())
                idx += 1
            idx += 1
            if lang == "text" and block and any("Android APP" in item for item in block):
                add_architecture_visual(doc)
            else:
                p = doc.add_paragraph()
                shade_paragraph(p, "F4F6F8")
                p.paragraph_format.left_indent = Pt(8)
                p.paragraph_format.right_indent = Pt(8)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = p.add_run("\n".join(block))
                set_run_font(run, name="Consolas", size=8.3, color=NAVY)
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[idx + 1].strip().lstrip("|")):
            table_lines = [stripped]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = [[strip_inline(cell.strip()) for cell in row.strip("|").split("|")] for row in table_lines]
            cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            if cols == 2:
                widths = [2600, 6760]
            elif cols == 3:
                widths = [2100, 3300, 3960]
            elif cols == 4:
                widths = [1500, 2520, 2520, 2820]
            else:
                base = 9360 // cols
                widths = [base] * cols
                widths[-1] += 9360 - sum(widths)
            set_table_width(table, widths, indent=120)
            for r_idx, row_data in enumerate(rows):
                set_table_row_cant_split(table.rows[r_idx])
                for c_idx in range(cols):
                    value = row_data[c_idx] if c_idx < len(row_data) else ""
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell, 95, 110, 95, 110)
                    set_cell_shading(cell, NAVY if r_idx == 0 else ("F7F9FB" if r_idx % 2 == 0 else WHITE))
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.05
                    add_rich_text(p, value)
                    for run in p.runs:
                        set_run_font(run, size=8.2 if cols >= 4 else 8.7, color=WHITE if r_idx == 0 else INK, bold=(r_idx == 0 or c_idx == 0))
                table.rows[r_idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            set_repeat_table_header(table.rows[0])
            if rows and rows[0] and "优先级" in rows[0]:
                for row_idx in range(1, len(rows)):
                    label = rows[row_idx][0] if rows[row_idx] else ""
                    fill = PALE_RED if label == "P0" else (PALE_GOLD if label == "P1" else PALE_BLUE)
                    set_cell_shading(table.cell(row_idx, 0), fill)
            doc.add_paragraph().paragraph_format.space_after = Pt(1)
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = strip_inline(heading_match.group(2))
            if level == 1:
                idx += 1
                continue
            style = "Heading 1" if level == 2 else ("Heading 2" if level == 3 else "Heading 3")
            p = doc.add_paragraph(title, style=style)
            if level == 2:
                set_paragraph_border(p, color=LINE, size=6, space=3)
            idx += 1
            continue
        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        number_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet_match or number_match:
            if bullet_match:
                p = doc.add_paragraph(style="List Bullet")
                add_rich_text(p, bullet_match.group(1))
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.32)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(3.5)
                p.paragraph_format.line_spacing = 1.15
                number_run = p.add_run(f"{number_match.group(1)}. ")
                set_run_font(number_run, size=10.2, color=INK)
                add_rich_text(p, number_match.group(2))
            set_keep_lines(p)
            idx += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Lead Callout")
            shade_paragraph(p, PALE_BLUE)
            if "已经完成 Android 健康管理 MVP" in stripped:
                set_page_break_before(p)
            add_rich_text(p, stripped.lstrip("> "))
            idx += 1
            continue
        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if not nxt or nxt.startswith(("#", "- ", ">", "```", "|")) or re.match(r"^\d+\.\s", nxt) or nxt == "---":
                break
            paragraph_lines.append(nxt)
            idx += 1
        text = " ".join(paragraph_lines)
        style = None
        if text.startswith("**当前边界**") or text.startswith("**尚未完成**") or text.startswith("**接口缺口提示**"):
            style = "Warning Callout"
        elif text.startswith("**MVP 发布状态"):
            style = "Warning Callout"
        p = doc.add_paragraph(style=style)
        if style:
            shade_paragraph(p, PALE_RED)
        add_rich_text(p, text)
        set_keep_lines(p)


def finalize_doc(doc: Document) -> None:
    for table in doc.tables:
        set_table_width(table, [cell._tc.tcPr.tcW.w for cell in table.rows[0].cells], indent=120)
        set_repeat_table_header(table.rows[0])
        for row in table.rows:
            set_table_row_cant_split(row)
    add_running_furniture(doc)
    props = doc.core_properties
    props.title = "睿禾精灵 APP 全景说明与技术架构"
    props.subject = "ReHealth AI 管理层产品与技术全景报告"
    props.author = "ReHealth AI / 睿禾健康"
    props.keywords = "ReHealth, 睿禾精灵, Android, 可穿戴, CVD, RHI, 技术架构"
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    for p in doc.paragraphs:
        if p.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            set_keep_with_next(p)
            set_keep_lines(p)
        if not p.text.strip():
            p.paragraph_format.space_after = Pt(0)


def build() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_business_summary(doc)
    parse_markdown(doc, source)
    finalize_doc(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
